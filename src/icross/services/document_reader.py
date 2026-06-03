"""Document reader service using Docling.

Provides unified document parsing for PDF, DOCX, XLSX, PPTX, HTML, images,
and TXT files. Returns Markdown-formatted text with table extraction and OCR.

Usage:
    from icross.services.document_reader import read_document

    result = read_document("/path/to/file.pdf")
    print(result["markdown"])
    print(result["metadata"])
"""

import json
from pathlib import Path
from typing import Any


def read_document(file_path: str) -> dict[str, Any]:
    """Parse a document file and return its content as Markdown.

    Args:
        file_path: Absolute path to the document file.

    Returns:
        Dict with:
            - success: bool
            - markdown: str (document content in Markdown format)
            - metadata: dict (file metadata)
            - pages: int (page count if applicable)
            - error: str (if success is False)
    """
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    ext = path.suffix.lower()

    try:
        return _try_docling(path)
    except ImportError:
        pass

    # Fallback: use specific libraries per file type
    if ext == ".pdf":
        return _fallback_pdf(path)
    elif ext in (".docx", ".doc"):
        return _fallback_docx(path)
    elif ext in (".xlsx", ".xls"):
        return _fallback_xlsx(path)
    elif ext == ".csv":
        return _fallback_csv(path)
    elif ext in (".txt", ".md", ".json", ".xml", ".yaml", ".yml", ".html", ".htm"):
        return _fallback_text(path)
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"):
        return _fallback_image(path)
    else:
        return {
            "success": False,
            "error": f"不支持的文件格式: {ext}",
            "hint": "请安装 Docling (pip install docling) 以获得完整格式支持",
        }


def _try_docling(path: Path) -> dict[str, Any]:
    """Parse document using Docling."""
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(path))

    doc = result.document

    # Export to markdown
    md_content = doc.export_to_markdown()

    # Basic metadata
    metadata = {}
    if hasattr(doc, "name") and doc.name:
        metadata["name"] = doc.name
    if hasattr(doc, "version") and doc.version:
        metadata["version"] = doc.version
    if hasattr(doc, "description") and doc.description:
        metadata["description"] = doc.description

    # Count pages if available
    pages = 0
    if hasattr(doc, "pages") and doc.pages:
        pages = len(doc.pages)

    # Collect references/citations
    refs = []
    if hasattr(doc, "references") and doc.references:
        for ref in doc.references:
            label = getattr(ref, "label", "")
            if label:
                refs.append(label)

    return {
        "success": True,
        "engine": "docling",
        "markdown": md_content,
        "metadata": metadata,
        "pages": pages,
        "references": refs,
        "file_name": path.name,
        "file_size_bytes": path.stat().st_size,
    }


def _fallback_pdf(path: Path) -> dict[str, Any]:
    """Fallback PDF reader."""
    try:
        # Try pypdf / PyPDF2
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError("No PDF library available")

        reader = PdfReader(str(path))
        pages_content = []
        for page in reader.pages:
            text = page.extract_text() if page else ""
            if text.strip():
                pages_content.append(text)

        return {
            "success": True,
            "engine": "pypdf",
            "markdown": "\n\n".join(pages_content),
            "metadata": {
                "pages": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None,
            },
            "pages": len(reader.pages),
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
        }
    except Exception as e:
        return {"success": False, "error": f"PDF 解析失败: {e}"}


def _fallback_docx(path: Path) -> dict[str, Any]:
    """Fallback DOCX reader using python-docx."""
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables_md = []
        for i, table in enumerate(doc.tables):
            rows_md = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_md.append(" | ".join(cells))
            if rows_md:
                tables_md.append(f"### 表格 {i+1}\n" + "\n".join(rows_md))

        content = "\n\n".join(paragraphs)
        if tables_md:
            content += "\n\n" + "\n\n".join(tables_md)

        return {
            "success": True,
            "engine": "python-docx",
            "markdown": content,
            "metadata": {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            },
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
        }
    except Exception as e:
        return {"success": False, "error": f"DOCX 解析失败: {e}"}


def _fallback_xlsx(path: Path) -> dict[str, Any]:
    """Fallback XLSX reader using openpyxl."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sheets_md = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_md = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rows_md.append(" | ".join(cells))

            sheet_content = f"## Sheet: {sheet_name}\n" + "\n".join(rows_md)
            sheets_md.append(sheet_content)

        wb.close()

        return {
            "success": True,
            "engine": "openpyxl",
            "markdown": "\n\n".join(sheets_md),
            "metadata": {
                "sheets": wb.sheetnames,
                "sheet_count": len(wb.sheetnames),
            },
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
        }
    except Exception as e:
        return {"success": False, "error": f"XLSX 解析失败: {e}"}


def _fallback_csv(path: Path) -> dict[str, Any]:
    """Fallback CSV reader."""
    try:
        import csv
        import io

        content = path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(io.StringIO(content))

        rows_md = []
        for row in reader:
            rows_md.append(" | ".join(row))

        return {
            "success": True,
            "engine": "csv",
            "markdown": "\n".join(rows_md),
            "metadata": {
                "rows": len(rows_md),
            },
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
        }
    except Exception as e:
        return {"success": False, "error": f"CSV 解析失败: {e}"}


def _fallback_text(path: Path) -> dict[str, Any]:
    """Read text-based files directly."""
    content = path.read_text(encoding="utf-8", errors="replace")
    return {
        "success": True,
        "engine": "text",
        "markdown": content,
        "metadata": {
            "characters": len(content),
            "lines": content.count("\n") + 1,
        },
        "file_name": path.name,
        "file_size_bytes": path.stat().st_size,
    }


def _fallback_image(path: Path) -> dict[str, Any]:
    """Fallback image reader — returns EXIF metadata and OCR hint."""
    metadata = {
        "format": path.suffix.lower().lstrip("."),
        "width": 0,
        "height": 0,
    }
    try:
        from PIL import Image
        img = Image.open(str(path))
        metadata["width"] = img.width
        metadata["height"] = img.height
        metadata["mode"] = img.mode
        img.close()
    except ImportError:
        metadata["note"] = "安装 Pillow 以获取图片尺寸"
    except Exception:
        pass

    return {
        "success": True,
        "engine": "pil",
        "markdown": f"![{path.name}]({path.name})\n\n图片尺寸: {metadata.get('width', '?')}x{metadata.get('height', '?')}",
        "metadata": metadata,
        "file_name": path.name,
        "file_size_bytes": path.stat().st_size,
        "ocr_hint": "需要安装 Docling 或 Tesseract 以提取图片中的文字",
    }
